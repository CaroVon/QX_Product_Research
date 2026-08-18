"""
============================================================
本地文件解析模块
—— PDF（PyMuPDF）/ TXT / Markdown / DOCX（可选 python-docx）
   每条切片以 local://{filename} 作为伪装 URL
============================================================
"""

import os

from app.rag.chunker import chunk_text


def parse_local_file(file_path: str, filename: str) -> list[dict]:
    """
    按扩展名分派解析器，提取全文文本，切片后返回结构化列表。

    Args:
        file_path: 文件磁盘绝对路径
        filename:  原始文件名（用于构造 local:// 伪装 URL）

    Returns:
        [{"content": "...", "url": "local://xxx.pdf"}, ...]

    Raises:
        ValueError: 不支持的文件类型 / 解析失败（调用方负责清理落盘文件）
    """
    ext = os.path.splitext(filename)[1].lower().lstrip(".") or ""
    if ext in ("pdf",):
        full_text = _parse_pdf(file_path)
    elif ext in ("txt", "md", "markdown"):
        full_text = _parse_plain_text(file_path)
    elif ext in ("docx", "doc"):
        full_text = _parse_docx(file_path)
    else:
        raise ValueError(f"不支持的文件类型: .{ext}")

    chunks = chunk_text(full_text)

    results: list[dict] = []
    for chunk in chunks:
        results.append({
            "content": chunk,
            "url": f"local://{filename}",
        })
    return results


def _parse_pdf(file_path: str) -> str:
    """PyMuPDF 提取 PDF 文本。"""
    try:
        import pymupdf  # 新版推荐入口（PyMuPDF 1.24+）
    except ImportError:  # pragma: no cover - 旧版兼容
        import fitz as pymupdf  # type: ignore[no-redef]

    doc = pymupdf.open(file_path)
    try:
        parts: list[str] = []
        for page in doc:
            text = page.get_text()
            if text:
                parts.append(text)
    finally:
        doc.close()
    return "\n".join(parts)


def _parse_plain_text(file_path: str) -> str:
    """TXT / Markdown 直接读取（容错编码）。"""
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    # 最后兜底：latin-1 永不失败（不丢数据）
    with open(file_path, "r", encoding="latin-1") as f:
        return f.read()


def _parse_docx(file_path: str) -> str:
    """DOCX 解析（依赖 python-docx）；.doc 旧格式无法直接解析。"""
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise ValueError(
            "DOCX 解析需要 python-docx，请执行: pip install python-docx"
        ) from exc

    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".doc":
        raise ValueError("旧版 .doc 格式暂不支持，请另存为 .docx 或 PDF 后上传")

    document = docx.Document(file_path)
    parts: list[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # 表格内容
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


# ─── 兼容旧调用方：parse_local_pdf（backend projects.py 等） ─────────
def parse_local_pdf(file_path: str, filename: str) -> list[dict]:
    """兼容旧签名：按扩展名自动分派解析（PDF/TXT/MD/DOCX）。"""
    return parse_local_file(file_path, filename)


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("用法: python local_parser.py <file>")
        sys.exit(1)
    path, name = sys.argv[1], os.path.basename(sys.argv[1])
    chunks = parse_local_file(path, name)
    print(f"共 {len(chunks)} 个切片")
    for c in chunks[:3]:
        print("----")
        print(c["content"][:200])
